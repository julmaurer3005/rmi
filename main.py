import os
import requests
import feedparser
import json
import re
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document

from cidades import CIDADES

USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
HEADERS = {"User-Agent": USER_AGENT}

def classificar(texto):
    """
    Classifica a notícia com base em palavras-chave no título/texto.
    """
    texto_lower = texto.lower()
    if any(p in texto_lower for p in ["bombeiro", "incêndio", "socorristas", "obito", "óbito", "acidente", "atropelamento"]):
        return "CBMRS"
    elif any(p in texto_lower for p in ["voluntário"]):
        return "Bombeiros Voluntários"
    elif any(p in texto_lower for p in ["busca", "resgate", "afogamento", "desaparecido", "resgatado", "resgatada"]):
        return "Buscas/Salvamento"
    elif any(p in texto_lower for p in ["chuva", "alagamento", "defesa civil", "deslizamento", "enchente", "temporal"]):
        return "Defesa Civil"
    elif any(p in texto_lower for p in ["desabamento", "enchurrada", "rio", "alagamento"]):
        return "Câmaras Municipais"
    else:
        return "Outros"

def limpar_texto_xml(texto):
    """
    Remove caracteres de controle invisíveis do texto coletado (como \x00, \x0b) 
    que fariam a biblioteca python-docx dar erro de XML na hora de salvar o documento.
    """
    if not isinstance(texto, str):
        return str(texto)
    # Remove caracteres inválidos em XML 1.0 (ver RFC 00-08, 0B-0C, 0E-1F)
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', texto)

def extrair_resumo_html(url):
    """
    Acessa o link da matéria e busca o primeiro parágrafo relevante para usar como resumo.
    """
    try:
        req = requests.get(url, headers=HEADERS, timeout=5)
        soup = BeautifulSoup(req.text, "html.parser")
        paragrafos = soup.find_all("p")
        
        for p in paragrafos:
            texto = p.get_text(strip=True)
            # Evita pegar links vazios ou parágrafos de créditos muito curtos
            if len(texto) > 60:
                # Limita o tamanho do resumo
                return texto[:300] + "..." if len(texto) > 300 else texto
                
    except Exception:
        return "Aviso: Não foi possível realizar o scraping automático deste conteúdo."
        
    return "Resumo não disponível na página."

def filtrar_por_cidade(titulo):
    """
    Usa Expressões Regulares com limites de palavra (\b) para combater o 
    falso positivo de substrings curtas, ignorando maiúsculas e minúsculas.
    """
    for cidade in CIDADES:
        padrao = r"(?i)\b" + re.escape(cidade) + r"\b"
        if re.search(padrao, titulo):
            return cidade
    return None

def gerar_relatorio(data_input):
    
    # Validações Iniciais
    if not os.path.exists("fontes.json"):
        print("Erro: Arquivo fontes.json não encontrado no diretório atual.")
        return
        
    with open("fontes.json", "r", encoding="utf-8") as f:
        fontes = json.load(f)

    noticias_processadas = []

    # 1. Coleta e Parsing de Dados
    for fonte in fontes:
        print(f"\n[+] Coletando da fonte: {fonte['nome']}...")
        try:
            if fonte["tipo"] == "rss":
                # Baixa o conteúdo primeiro com requests para forçar um timeout e evitar travamentos
                resposta_rss = requests.get(fonte["url"], headers=HEADERS, timeout=10)
                feed = feedparser.parse(resposta_rss.content)
                # Pega as últimas 20 notícias do RSS
                for entry in feed.entries[:20]:
                    # Limpeza de HTML usando BeautifulSoup caso o resumo venha sujo
                    raw_summary = getattr(entry, "summary", "Sem resumo.")
                    clean_summary = BeautifulSoup(raw_summary, "html.parser").get_text(separator=' ', strip=True) if raw_summary else "Sem resumo."
                    
                    noticias_processadas.append({
                        "titulo": entry.title,
                        "link": entry.link,
                        "resumo": clean_summary[:300] + ("..." if len(clean_summary) > 300 else ""),
                        "fonte": fonte["nome"]
                    })
                    
            elif fonte["tipo"] == "site":
                response = requests.get(fonte["url"], headers=HEADERS, timeout=10)
                soup = BeautifulSoup(response.text, "html.parser")
                
                # Procura por ambos H3 ou H2 no HTML e links directos grandes
                itens = soup.select("h3 a, h2 a")
                if not itens:
                    # Fallback para sites mal estruturados
                    itens = [i for i in soup.select("a") if len(i.get_text(strip=True)) > 30]

                count = 0
                for item in itens:
                    if count >= 10: # Limite por portal HTML para não demorar
                        break
                        
                    titulo = item.get_text(strip=True)
                    link = item.get("href")
                    
                    if not link or not titulo:
                        continue
                        
                    # Corrige URLs relativas (ex: /noticia/1234 -> https://site.com/noticia/1234)
                    if link.startswith("/"):
                        if fonte["url"].endswith("/"):
                            link = fonte["url"][:-1] + link
                        else:
                            # Tenta remontar o hostname
                            domain_parts = fonte["url"].split("/")
                            base_url = domain_parts[0] + "//" + domain_parts[2]
                            link = base_url + link
                        
                    print(f"  - Extraindo resumo: {titulo[:40]}...")
                    resumo = extrair_resumo_html(link)
                    
                    noticias_processadas.append({
                        "titulo": titulo,
                        "link": link,
                        "resumo": resumo,
                        "fonte": fonte["nome"]
                    })
                    count += 1
                    
        except Exception as e:
            print(f"[!] Erro na fonte {fonte['nome']}: {e}")

    # 2. Avaliação e Classificação Semântica
    relatorio = {
        "CBMRS": [],
        "bombeiros": [],
        "bombeiro": [],
        "Bombeiros Voluntários": [],
        "Buscas/Salvamento": [],
        "Defesa Civil": [],
        "Câmaras Municipais": [],
        "Outros": []
    }

    print("\n[*] Filtrando citações por município e classificando...")
    for n in noticias_processadas:
        # Busca a cidade tanto no título quanto no resumo (1º parágrafo da notícia)
        texto_busca = n["titulo"] + " " + n["resumo"]
        cidade_mencionada = filtrar_por_cidade(texto_busca)
        
        # Notícias apenas entram se citarem explicitamente um dos 47 municípios alvo
        if cidade_mencionada:
            n["cidade"] = cidade_mencionada
            categoria = classificar(n["titulo"])
            relatorio[categoria].append(n)

    # 3. Empacotamento / Output
    print("\n[+] Agrupando e gerando o documento Word...")
    os.makedirs("output", exist_ok=True)
    
    doc = Document()
    doc.add_heading("RELATÓRIO DE MONITORAMENTO REGIONAL", 0)
    doc.add_paragraph(f"Data Referência: {data_input}")

    total_sucesso = 0
    for categoria, itens in relatorio.items():
        doc.add_heading(categoria, level=1)
        
        if not itens:
            doc.add_paragraph("· Sem novidades reportadas para este segmento.")
        else:
            for n in itens:
                total_sucesso += 1
                doc.add_paragraph(f"· Município: {limpar_texto_xml(n['cidade'])}")
                doc.add_paragraph(f"Matéria: {limpar_texto_xml(n['titulo'])}")
                doc.add_paragraph(f"Resumo: {limpar_texto_xml(n['resumo'])}")
                doc.add_paragraph(f"Acesso via: {limpar_texto_xml(n['link'])} (Fonte: {n['fonte']})")
                doc.add_paragraph("")

    # Gerando o arquivo final com timestamp para não sobrecrever arquivos passados
    timestamp_formatado = datetime.now().strftime("%Y-%m-%d_%H-%M")
    nome_arquivo = f"output/Relatorio_Inteligencia_{timestamp_formatado}.docx"
    doc.save(nome_arquivo)

    print(f"\n[OK] Finalizado! {total_sucesso} notícia(s) relevantes foram computadas e filtradas positivamente.")
    print(f"     O seu Word foi salvo em: {nome_arquivo}\n")
    return nome_arquivo, total_sucesso

def main():
    print("="*40)
    print(" GERADOR DE RELATÓRIO DE INTELIGÊNCIA ")
    print("="*40)
    data_input = input("Digite a data para exibir no relatório (Ex: 10/04/2026): ")
    gerar_relatorio(data_input)

if __name__ == "__main__":
    main()